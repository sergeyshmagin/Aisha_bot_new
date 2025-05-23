"""
Сервис для обработки текста и генерации документов
"""
import io
import logging
from typing import Dict, List, Optional

import aiohttp
from docx import Document
from docx.shared import Pt, RGBColor

from app.prompts.transcribe.prompts import (
    SHORT_SUMMARY_PROMPT,
    MOM_PROMPT,
    TODO_PROMPT,
    PROTOCOL_PROMPT,
    FULL_TRANSCRIPT_PROMPT,
)

from app.core.config import settings
from app.services.base import BaseService
from app.shared.utils.openai import get_openai_headers

logger = logging.getLogger(__name__)


class TextProcessingService(BaseService):
    """
    Сервис для обработки текста и генерации документов
    """

    async def generate_summary(self, text: str, max_length: int = 2000) -> str:
        """
        Генерирует краткое содержание текста
        
        Args:
            text: Исходный текст
            max_length: Максимальная длина краткого содержания
            
        Returns:
            str: Краткое содержание
        """
        prompt = f"{SHORT_SUMMARY_PROMPT}\n\nТекст:\n{text[:max_length]}"
        
        return await self._process_with_gpt(prompt)
    
    async def generate_todo_list(self, text: str, max_length: int = 2000) -> List[str]:
        """
        Генерирует список задач из текста
        
        Args:
            text: Исходный текст
            max_length: Максимальная длина обрабатываемого текста
            
        Returns:
            List[str]: Список задач
        """
        prompt = f"{TODO_PROMPT}\n\nТекст:\n{text[:max_length]}"
        
        result = await self._process_with_gpt(prompt)
        
        # Разбиваем на отдельные задачи
        tasks = []
        for line in result.split("\n"):
            line = line.strip()
            if line.startswith("-") or line.startswith("•"):
                tasks.append(line[1:].strip())
            elif line and not any(task.startswith(line) for task in tasks):
                tasks.append(line)
        
        return tasks
    
    async def generate_meeting_protocol(self, text: str, metadata: Optional[Dict] = None) -> bytes:
        """
        Генерирует протокол встречи в формате Word
        
        Args:
            text: Исходный текст
            metadata: Метаданные встречи (название, дата, участники)
            
        Returns:
            bytes: Документ в формате Word
        """
        metadata = metadata or {}
        
        # Получаем структурированный протокол через GPT
        prompt = f"{MOM_PROMPT}\n\nТранскрипция:\n{text[:4000]}"
        
        protocol_text = await self._process_with_gpt(prompt)
        
        # Создаем документ Word
        doc = Document()
        
        # Заголовок
        title = doc.add_heading("ПРОТОКОЛ ВСТРЕЧИ", level=1)
        title.alignment = 1  # По центру
        
        # Метаданные
        if metadata.get("title"):
            doc.add_paragraph(f"Название: {metadata['title']}")
        if metadata.get("date"):
            doc.add_paragraph(f"Дата: {metadata['date']}")
        if metadata.get("participants"):
            doc.add_paragraph(f"Участники: {metadata['participants']}")
        
        # Добавляем содержимое протокола
        for line in protocol_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            
            # Определяем уровень заголовка
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                # Маркированный список
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("1. ") or line.startswith("1) "):
                # Нумерованный список
                doc.add_paragraph(line[3:], style="List Number")
            else:
                # Обычный текст
                doc.add_paragraph(line)
        
        # Сохраняем документ в память
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.read()
    
    async def _process_with_gpt(self, prompt: str) -> str:
        """
        Обрабатывает текст с помощью GPT
        
        Args:
            prompt: Промпт для GPT
            
        Returns:
            str: Результат обработки
        """
        url = "https://api.openai.com/v1/chat/completions"
        headers = get_openai_headers(settings.OPENAI_API_KEY)
        data = {
            "model": "gpt-4o",
            "messages": [
                {"role": "system", "content": "Ты помощник, который обрабатывает текст и создает структурированные документы."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 1500
        }
        
        # Если API ключ не указан, возвращаем заглушку
        if not settings.OPENAI_API_KEY:
            logger.warning("Отсутствует API ключ OpenAI, возвращаем заглушку")
            return "Для обработки текста необходимо настроить API ключ OpenAI."
        
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(url, headers=headers, json=data) as response:
                    if response.status == 200:
                        result = await response.json()
                        return result["choices"][0]["message"]["content"]
                    else:
                        error_text = await response.text()
                        logger.error(f"GPT API error: {error_text}")
                        return "Ошибка при обработке текста."
        except Exception as e:
            logger.error(f"Error calling GPT API: {e}")
            return "Ошибка при обработке текста."

    async def process_text(self, text: str) -> str:
        """
        Обрабатывает текст (основной метод для TranscriptProcessingHandler)
        
        Args:
            text: Исходный текст
            
        Returns:
            str: Обработанный текст
        """
        try:
            # Очищаем текст от лишних пробелов и переносов строк
            text = "\n".join(line.strip() for line in text.split("\n") if line.strip())
            
            # Удаляем множественные пробелы
            text = " ".join(text.split())
            
            # Добавляем переносы строк после точек
            text = text.replace(". ", ".\n")
            
            # Форматируем абзацы
            paragraphs = text.split("\n\n")
            formatted_paragraphs = []
            for p in paragraphs:
                if p.strip():
                    # Добавляем отступ в начале абзаца
                    formatted_paragraphs.append("    " + p.strip())
            
            # Собираем текст обратно
            processed_text = "\n\n".join(formatted_paragraphs)
            
            logger.info(f"[TEXT] Текст обработан, длина: {len(processed_text)}")
            return processed_text
            
        except Exception as e:
            logger.error(f"[TEXT] Ошибка при обработке текста: {e}")
            return text  # В случае ошибки возвращаем исходный текст
    
    async def format_summary(self, text: str) -> str:
        """
        Создает краткое содержание из транскрипта
        
        Args:
            text: Исходный текст транскрипта
            
        Returns:
            str: Краткое содержание
        """
        return await self.generate_summary(text)
    
    async def format_todo(self, text: str) -> str:
        """
        Создает список задач из транскрипта
        
        Args:
            text: Исходный текст транскрипта
            
        Returns:
            str: Список задач в текстовом формате
        """
        tasks = await self.generate_todo_list(text)
        return "\n".join([f"- {task}" for task in tasks])
    
    async def format_protocol(self, text: str) -> str:
        """
        Создает протокол встречи из транскрипта
        
        Args:
            text: Исходный текст транскрипта
            
        Returns:
            str: Протокол в текстовом формате
        """
        prompt = f"{PROTOCOL_PROMPT}\n\nТранскрипция:\n{text[:4000]}"
        return await self._process_with_gpt(prompt)
